"""
Parsers for Excel, Word, and PDF files from Google Drive.
Handles real NeGD file formats discovered via full Drive scan.
"""

import os
import re
from datetime import datetime

import pandas as pd
import openpyxl
import docx
import pdfplumber


# -------------------------------------------------------------------------
# Excel parser — with smart header detection
# -------------------------------------------------------------------------
def parse_excel(file_path):
    """Parse an Excel file and return all sheets as a dict of DataFrames."""
    try:
        xls = pd.ExcelFile(file_path, engine="openpyxl")
        sheets = {}
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(xls, sheet_name=sheet_name)
            if not df.empty:
                sheets[sheet_name] = df
        return sheets
    except Exception as e:
        return {"error": str(e)}


def _detect_header_row(df, keywords):
    """Detect if the real header is in one of the first few data rows."""
    for row_idx in range(min(5, len(df))):
        row_values = [str(v).lower().strip() for v in df.iloc[row_idx].values if pd.notna(v)]
        matches = sum(1 for val in row_values for kw in keywords if kw in val)
        if matches >= 2:
            return row_idx
    return None


def _read_sheet_smart(file_path, sheet_name, hint_keywords):
    """Read a sheet with smart header detection."""
    try:
        df = pd.read_excel(file_path, sheet_name=sheet_name, engine="openpyxl")
        if df.empty:
            return df
        header_row = _detect_header_row(df, hint_keywords)
        if header_row is not None:
            df = pd.read_excel(file_path, sheet_name=sheet_name, header=header_row + 1, engine="openpyxl")
        df.columns = [str(c).strip() for c in df.columns]
        df = df.dropna(how="all")
        return df
    except Exception:
        return pd.DataFrame()


SKIP_NAMES = {"name", "resource", "employee", "person", "emp no.", "sr no", "s.no",
              "sr no.", "sl. no.", "sl.no", "s. no.", "task name", "task", "category id",
              "team member", "team member's", "#", "sn"}

SKIP_ROLES = {"junior developer", "quality engineer", "production support trainee",
              "senior developer", "devops engineer", "qa engineer", "developer",
              "project manager", "business analyst", "tech lead"}


# -------------------------------------------------------------------------
# TASK EXTRACTION
# -------------------------------------------------------------------------
def extract_tasks_from_excel(file_path, project_name):
    """Extract tasks. Handles:
    - Standard: Task + Status + Owner + Date columns
    - OpenForge: aid + title + status + assigned_to + end_date
    - Daily Tracker: Date + Team Member + Module + Priority
    - Sprint plans with task rows
    """
    sheets = parse_excel(file_path)
    if "error" in sheets:
        return pd.DataFrame()

    task_hints = ["task", "activity", "description", "title", "item", "sr.no", "s.no",
                  "sub-task", "deliverable", "action", "aid", "issue", "bug", "checklist"]
    status_hints = ["status", "state", "progress", "completion"]

    all_tasks = []
    xls = pd.ExcelFile(file_path, engine="openpyxl")

    for sheet_name in xls.sheet_names:
        # Try direct read first, then smart read if columns look bad
        df = pd.read_excel(xls, sheet_name=sheet_name)
        if df.empty or len(df) < 1:
            continue
        df.columns = [str(c).strip() for c in df.columns]

        # Check if headers look like data (too many Unnamed columns)
        unnamed_count = sum(1 for c in df.columns if "unnamed" in str(c).lower())
        if unnamed_count > len(df.columns) * 0.5:
            df = _read_sheet_smart(file_path, sheet_name, task_hints + status_hints)
            if df.empty:
                continue

        cols_lower = {c: c.lower().strip() for c in df.columns}
        df_r = df.rename(columns={c: cols_lower[c] for c in df.columns})

        # --- OpenForge artifact format (aid, title, status, assigned_to) ---
        if "aid" in df_r.columns and "title" in df_r.columns:
            for _, row in df_r.iterrows():
                title = _safe_str(row.get("title") or row.get("task_title", ""))
                if not title or len(title) < 3:
                    continue
                status = _normalize_status(_safe_str(row.get("status", "Unknown")))
                assigned = _safe_str(row.get("assigned_to", ""))
                date_val = _safe_date(row.get("actual_end_date") or row.get("end_date", ""))
                priority = _normalize_priority(_safe_str(row.get("priority", "Medium")))
                all_tasks.append(_task_row(project_name, title, status, assigned, date_val, priority))
            continue

        # --- Standard task columns ---
        task_col = _find_col(df_r, ["task name", "task title", "task", "activity",
                                     "task description", "description", "sub-task",
                                     "action items", "change description", "task lists",
                                     "issue", "task definition", "title", "item",
                                     "checklist items", "action items (mvp)",
                                     "task", "module"])
        status_col = _find_col(df_r, ["status", "state", "progress", "completion",
                                       "released to prod"])
        assigned_col = _find_col(df_r, ["task owner", "task owner\n(negd)",
                                         "assigned to", "assigned_to", "owner",
                                         "ownership", "responsible",
                                         "team member name", "team member's",
                                         "resource name"])
        date_col = _find_col(df_r, ["end date", "actual completion date",
                                     "completion date", "closing week",
                                     "due date", "deadline", "expected completion date",
                                     "tentative date", "completed date",
                                     "end\ndate", "key date"])
        priority_col = _find_col(df_r, ["priority", "priority (h/m/l)",
                                         "priority (p1/p2/p3)", "severity"])

        if not task_col:
            continue

        for _, row in df_r.iterrows():
            task_name = _safe_str(row.get(task_col, ""))
            if not task_name or len(task_name) < 3:
                continue
            if task_name.lower() in SKIP_NAMES:
                continue

            status = _normalize_status(_safe_str(row.get(status_col, "Unknown")) if status_col else "Unknown")
            assigned = _safe_str(row.get(assigned_col, "")) if assigned_col else ""
            date_val = _safe_date(row.get(date_col, "")) if date_col else ""
            priority = _normalize_priority(_safe_str(row.get(priority_col, "Medium")) if priority_col else "Medium")

            all_tasks.append(_task_row(project_name, task_name, status, assigned, date_val, priority))

    return pd.DataFrame(all_tasks) if all_tasks else pd.DataFrame()


def _task_row(project, name, status, assigned, date_val, priority):
    return {
        "project": project,
        "task_name": name[:120],
        "status": status,
        "assigned_to": assigned if assigned.lower() not in SKIP_NAMES else "",
        "closed_date": date_val if status == "Closed" else None,
        "priority": priority,
    }


# -------------------------------------------------------------------------
# RESOURCE EXTRACTION
# -------------------------------------------------------------------------
def extract_resources_from_excel(file_path, project_name):
    """Extract resources. Handles:
    - Timesheet tabs: "Amit (Dev)", "Madhur (PM)"
    - Structured tables: Name + Role/Designation + Agency
    - HR Summary: Role + Count + Company
    - Team Details: Team Member + Role + Email
    """
    sheets = parse_excel(file_path)
    if "error" in sheets:
        return pd.DataFrame()

    all_resources = []
    seen_names = set()

    # Strategy 1: Person-named tabs (timesheet format)
    person_tab = re.compile(r'^(.+?)\s*\((.+?)\)\s*$')
    for sheet_name in sheets.keys():
        match = person_tab.match(sheet_name)
        if match:
            name = match.group(1).strip()
            role_hint = match.group(2).strip()
            if name.lower() not in seen_names and len(name) > 1 and not name.replace(".", "").isdigit():
                seen_names.add(name.lower())
                role = _expand_role(role_hint)
                all_resources.append(_resource_row(name, role, project_name))

    # Strategy 2: Structured resource tables
    resource_keywords = ["name", "team member", "resource", "employee", "emp",
                         "role", "designation", "position", "agency"]
    xls = pd.ExcelFile(file_path, engine="openpyxl")

    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        if df.empty:
            continue
        df.columns = [str(c).strip() for c in df.columns]
        unnamed_count = sum(1 for c in df.columns if "unnamed" in str(c).lower())
        if unnamed_count > len(df.columns) * 0.5:
            df = _read_sheet_smart(file_path, sheet_name, resource_keywords)
            if df.empty:
                continue

        cols_lower = {c: c.lower().strip() for c in df.columns}
        df_r = df.rename(columns={c: cols_lower[c] for c in df.columns})

        # Look for Name + Role columns
        name_col = _find_col(df_r, ["name", "team member", "emp name", "employee name",
                                     "resource name"])
        role_col = _find_col(df_r, ["role", "designation", "designation name",
                                     "functional role", "position", "title"])

        if name_col and role_col:
            for _, row in df_r.iterrows():
                name = _safe_str(row.get(name_col, ""))
                if not _valid_person_name(name) or name.lower() in seen_names:
                    continue
                seen_names.add(name.lower())
                role = _safe_str(row.get(role_col, ""))
                all_resources.append(_resource_row(name, role, project_name))
        elif name_col:
            # Name column only (no role), still useful
            for _, row in df_r.iterrows():
                name = _safe_str(row.get(name_col, ""))
                if not _valid_person_name(name) or name.lower() in seen_names:
                    continue
                seen_names.add(name.lower())
                all_resources.append(_resource_row(name, "", project_name))

    # Strategy 3: HR Summary format (Role + Count — aggregate, not individual)
    for sheet_name in xls.sheet_names:
        df = _read_sheet_smart(file_path, sheet_name, ["role", "count", "resources"])
        if df.empty:
            continue
        cols_lower = {c: c.lower().strip() for c in df.columns}
        df_r = df.rename(columns={c: cols_lower[c] for c in df.columns})
        role_col = _find_col(df_r, ["role"])
        count_col = _find_col(df_r, ["number of resources", "count", "no. of resources",
                                      "current number of resources"])
        company_col = _find_col(df_r, ["company name", "agency"])
        if role_col and count_col:
            for _, row in df_r.iterrows():
                role = _safe_str(row.get(role_col, ""))
                if not role or role.lower() in SKIP_NAMES or role == "nan":
                    continue
                try:
                    count = int(float(row.get(count_col, 0)))
                except (ValueError, TypeError):
                    count = 0
                company = _safe_str(row.get(company_col, "")) if company_col else ""
                for i in range(min(count, 50)):
                    rname = f"{role} #{i+1}" + (f" ({company})" if company else "")
                    if rname.lower() not in seen_names:
                        seen_names.add(rname.lower())
                        all_resources.append(_resource_row(rname, role, project_name))

    return pd.DataFrame(all_resources) if all_resources else pd.DataFrame()


def _resource_row(name, role, project):
    bucket = _classify_role(role)
    return {
        "name": name,
        "role": role if role != "nan" else "",
        "bucket": bucket,
        "project": project,
        "tasks_completed_15d": 0,
        "tasks_pending": 0,
    }


def _valid_person_name(name):
    """Check if a string looks like a valid person name."""
    if not name or name == "nan" or len(name) < 2:
        return False
    if name.lower() in SKIP_NAMES | SKIP_ROLES:
        return False
    if name.replace(".", "").replace("-", "").replace(" ", "").isdigit():
        return False
    return True


# -------------------------------------------------------------------------
# FINANCIAL EXTRACTION
# -------------------------------------------------------------------------
def extract_financials_from_excel(file_path, project_name):
    """Extract financials. Handles:
    - Budget/Cost sheets: Role + Cost + Count
    - Subscription/billing: Tool + Last Billed Amount
    - Infra cost: Environment + Cost
    """
    sheets = parse_excel(file_path)
    if "error" in sheets:
        return pd.DataFrame()

    for sheet_name, df in sheets.items():
        cols_lower = {c: c.lower().strip() for c in df.columns}
        df_r = df.rename(columns={c: cols_lower[c] for c in df.columns})

        # Pattern 1: Budget/contracted + utilised
        budget_col = _find_col(df_r, ["budget", "contracted", "sanctioned", "total",
                                       "total monthly cost", "4 month cost",
                                       "12 months cost", "amount", "value"])
        utilised_col = _find_col(df_r, ["utilised", "utilized", "spent",
                                         "expenditure", "consumed", "last billed amount"])

        if budget_col:
            try:
                contracted = pd.to_numeric(df_r[budget_col], errors="coerce").sum()
                utilised = pd.to_numeric(df_r[utilised_col], errors="coerce").sum() if utilised_col else 0
                if contracted > 0:
                    if contracted > 10000000:
                        contracted /= 10000000
                        utilised /= 10000000
                    elif contracted > 100000:
                        contracted /= 10000000
                        utilised /= 10000000
                    pct = round(utilised / contracted * 100, 1) if contracted > 0 else 0
                    health = "Critical" if pct > 90 else ("Monitor" if pct > 75 else "Healthy")
                    return pd.DataFrame([{
                        "project": project_name,
                        "contracted_cr": round(contracted, 2),
                        "utilised_cr": round(utilised, 2),
                        "utilised_pct": pct,
                        "remaining_cr": round(contracted - utilised, 2),
                        "health": health,
                    }])
            except Exception:
                pass

    return pd.DataFrame()


# -------------------------------------------------------------------------
# RISK EXTRACTION (NEW)
# -------------------------------------------------------------------------
def extract_risks_from_excel(file_path, project_name):
    """Extract risks/issues. Handles:
    - Key Issues & Challenges: Issue + Status + Expected Date
    - Risk Register: Risk + Impact + Probability
    - Bug Reports: Issue + Severity + Status
    """
    sheets = parse_excel(file_path)
    if "error" in sheets:
        return pd.DataFrame()

    risk_hints = ["issue", "risk", "bug", "challenge", "problem", "severity",
                  "impact", "probability"]
    all_risks = []
    xls = pd.ExcelFile(file_path, engine="openpyxl")

    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        if df.empty or len(df) < 1:
            continue
        df.columns = [str(c).strip() for c in df.columns]
        unnamed_count = sum(1 for c in df.columns if "unnamed" in str(c).lower())
        if unnamed_count > len(df.columns) * 0.5:
            df = _read_sheet_smart(file_path, sheet_name, risk_hints)
            if df.empty:
                continue

        cols_lower = {c: c.lower().strip() for c in df.columns}
        df_r = df.rename(columns={c: cols_lower[c] for c in df.columns})

        desc_col = _find_col(df_r, ["issue", "risk", "bug", "challenge",
                                     "problem", "description", "risk description",
                                     "change description", "enhancements"])
        severity_col = _find_col(df_r, ["severity", "impact", "priority",
                                         "priority (h/m/l)", "risk level"])
        status_col = _find_col(df_r, ["status", "state", "track"])
        owner_col = _find_col(df_r, ["owner", "assigned to", "responsible",
                                      "task owner", "ownership"])
        date_col = _find_col(df_r, ["expected completion date", "due date",
                                     "end date", "deadline", "target date"])

        if not desc_col:
            continue

        for _, row in df_r.iterrows():
            desc = _safe_str(row.get(desc_col, ""))
            if not desc or len(desc) < 5 or desc.lower() in SKIP_NAMES:
                continue

            severity = _safe_str(row.get(severity_col, "Medium")) if severity_col else "Medium"
            sev_lower = severity.lower()
            if any(w in sev_lower for w in ["high", "h", "critical", "p1"]):
                severity = "High"
            elif any(w in sev_lower for w in ["low", "l", "p3", "minor"]):
                severity = "Low"
            else:
                severity = "Medium"

            owner = _safe_str(row.get(owner_col, "")) if owner_col else ""
            due = _safe_date(row.get(date_col, "")) if date_col else ""

            all_risks.append({
                "project": project_name,
                "severity": severity,
                "description": desc[:200],
                "owner": owner,
                "due_date": due,
            })

    return pd.DataFrame(all_risks) if all_risks else pd.DataFrame()


# -------------------------------------------------------------------------
# MEETING EXTRACTION FROM EXCEL (NEW)
# -------------------------------------------------------------------------
def extract_meetings_from_excel(file_path, project_name):
    """Extract meeting dates from structured ROD spreadsheets.
    Handles DPIIT 'ROD- Last Meeting with Client' format.
    """
    sheets = parse_excel(file_path)
    if "error" in sheets:
        return []

    meetings = []
    meeting_hints = ["meeting date", "date", "meeting", "rod"]

    xls = pd.ExcelFile(file_path, engine="openpyxl")
    for sheet_name in xls.sheet_names:
        df = _read_sheet_smart(file_path, sheet_name, meeting_hints)
        if df.empty:
            continue
        cols_lower = {c: c.lower().strip() for c in df.columns}
        df_r = df.rename(columns={c: cols_lower[c] for c in df.columns})

        date_col = _find_col(df_r, ["meeting date", "date of consultation", "date"])
        if date_col:
            for _, row in df_r.iterrows():
                date_val = _safe_date(row.get(date_col, ""))
                if date_val:
                    meetings.append({
                        "project": project_name,
                        "meeting_date": date_val,
                    })
    return meetings


# -------------------------------------------------------------------------
# Word parser
# -------------------------------------------------------------------------
def parse_word(file_path):
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
            if rows:
                tables.append(rows)
        return {"paragraphs": paragraphs, "tables": tables}
    except Exception as e:
        return {"error": str(e)}


def extract_meeting_from_word(file_path, project_name):
    content = parse_word(file_path)
    if "error" in content:
        return {}
    full_text = " ".join(content.get("paragraphs", []))
    meeting_date = _extract_date(full_text)
    return {"project": project_name, "meeting_date": meeting_date} if meeting_date else {}


# -------------------------------------------------------------------------
# PDF parser
# -------------------------------------------------------------------------
def parse_pdf(file_path):
    try:
        text_pages = []
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages[:5]:  # limit to 5 pages for speed
                text = page.extract_text()
                if text:
                    text_pages.append(text)
        return {"full_text": "\n".join(text_pages)}
    except Exception as e:
        return {"error": str(e)}


def extract_meeting_from_pdf(file_path, project_name):
    content = parse_pdf(file_path)
    if "error" in content:
        return {}
    meeting_date = _extract_date(content.get("full_text", ""))
    return {"project": project_name, "meeting_date": meeting_date} if meeting_date else {}


# -------------------------------------------------------------------------
# Helpers
# -------------------------------------------------------------------------
def _find_col(df, keywords):
    """Find column by keyword. Prioritizes exact match, then contains."""
    cols = list(df.columns)
    for col in cols:
        cl = str(col).lower().strip()
        for kw in keywords:
            if cl == kw:
                return col
    for col in cols:
        cl = str(col).lower().strip()
        for kw in keywords:
            if kw in cl:
                return col
    return None


def _safe_str(val):
    if pd.isna(val):
        return ""
    s = str(val).strip()
    return "" if s == "nan" else s


def _safe_date(val):
    if pd.isna(val) or val == "" or val is None:
        return ""
    s = str(val).strip()
    if s == "nan":
        return ""
    return s[:10]


def _normalize_status(s):
    sl = s.lower()
    if any(w in sl for w in ["done", "complete", "closed", "resolved", "delivered", "y", "completed"]):
        return "Closed"
    if any(w in sl for w in ["progress", "ongoing", "wip", "in-progress", "started", "inprogress"]):
        return "In Progress"
    if any(w in sl for w in ["pending", "open", "not started", "to do", "todo", "planned", "blocked"]):
        return "Pending"
    if not s or s == "Unknown":
        return "Unknown"
    return s


def _normalize_priority(s):
    sl = s.lower()
    if any(w in sl for w in ["p1", "high", "critical", "h"]):
        return "High"
    if any(w in sl for w in ["p3", "low", "l"]):
        return "Low"
    return "Medium"


def _classify_role(role):
    rl = str(role).lower()
    pm_kw = ["project manager", "pm", "business analyst", "ba", "scrum",
             "lead", "coordinator", "pmo", "manager", "analyst", "consultant",
             "program", "product"]
    return "PM" if any(kw in rl for kw in pm_kw) else "Tech"


def _expand_role(hint):
    mapping = {
        "dev": "Developer", "pm": "Project Manager", "qa": "QA Engineer",
        "uiux": "UI/UX Designer", "ui": "UI Designer", "ux": "UX Designer",
        "ba": "Business Analyst", "devops": "DevOps Engineer", "tpm": "Technical PM",
        "lead": "Tech Lead", "fe": "Frontend Developer", "be": "Backend Developer",
        "test": "QA Engineer", "design": "Designer", "scrum": "Scrum Master",
    }
    return mapping.get(hint.lower().strip(), hint)


def _extract_date(text):
    patterns = [
        r"(\d{1,2}[\s/-]\w{3,9}[\s/-]\d{4})",
        r"(\d{1,2}[/-]\d{1,2}[/-]\d{4})",
        r"(\w{3,9}\s+\d{1,2},?\s+\d{4})",
        r"[Dd]ated?\s*[:.]?\s*(\d{1,2}[.\-/]\d{1,2}[.\-/]\d{2,4})",
    ]
    for p in patterns:
        m = re.search(p, text)
        if m:
            return m.group(1)
    return None
